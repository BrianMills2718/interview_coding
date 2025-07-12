"""Executive brief generator for o3 methodology.

Creates semi-automated executive brief using python-docx as specified by ChatGPT-o3.
"""

import pandas as pd
import json
from pathlib import Path
import os
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = Path(os.getenv("OUTPUTS_DIR", ROOT / "outputs"))
PROC_DIR = Path(os.getenv("PROCESSED_DIR", ROOT / "data" / "processed"))


def load_theme_matrix(transcript_id: str) -> pd.DataFrame:
    """Load theme matrix for prevalence data."""
    matrix_path = OUTPUT_DIR / f"{transcript_id}_theme_matrix.xlsx"
    if matrix_path.exists():
        return pd.read_excel(matrix_path, sheet_name='Theme_Matrix', index_col=0)
    return pd.DataFrame()


def load_alpha_report(transcript_id: str) -> dict:
    """Load alpha reliability report."""
    alpha_path = OUTPUT_DIR / f"{transcript_id}_alpha.json"
    if alpha_path.exists():
        with open(alpha_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}


def load_merged_tags(transcript_id: str) -> pd.DataFrame:
    """Load merged tags for exemplar quotes."""
    merged_path = OUTPUT_DIR / f"{transcript_id}_tags_merged.csv"
    if merged_path.exists():
        return pd.read_csv(merged_path)
    return pd.DataFrame()


def get_exemplar_quotes(transcript_id: str, theme: str) -> list:
    """Get exemplar quotes for a theme by looking up UIDs."""
    # Load merged tags to find UIDs with the theme
    merged_df = load_merged_tags(transcript_id)
    if merged_df.empty:
        return []
    
    # Find UIDs where this theme is present (value = 1.0)
    if theme in merged_df.columns:
        theme_uids = merged_df[merged_df[theme] == 1.0]['uid'].tolist()
    else:
        return []
    
    # Load original tidy CSV to get the actual text
    tidy_path = PROC_DIR / f"{transcript_id}_tidy.csv"
    if not tidy_path.exists():
        return []
    
    tidy_df = pd.read_csv(tidy_path)
    
    # Get quotes for the theme UIDs
    quotes = []
    for uid in theme_uids[:3]:  # Limit to 3 exemplars
        row = tidy_df[tidy_df['uid'] == uid]
        if not row.empty:
            speaker = row.iloc[0]['speaker']
            text = row.iloc[0]['text']
            quotes.append(f'"{text}" - {speaker}')
    
    return quotes


def create_executive_brief(transcript_id: str) -> None:
    """Create semi-automated executive brief."""
    print(f"Creating executive brief for {transcript_id}...")
    
    # Load data
    theme_matrix = load_theme_matrix(transcript_id)
    alpha_report = load_alpha_report(transcript_id)
    
    if theme_matrix.empty:
        print(f"No theme matrix found for {transcript_id}")
        return
    
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading('AI Needs & Barriers – RAND Focus Groups (v1.0)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Top-line Findings
    doc.add_heading('1. Top-line Findings', level=1)
    
    # Get themes with prevalence > 70%
    high_prevalence = theme_matrix[theme_matrix['Prevalence_%'] > 70]
    
    if not high_prevalence.empty:
        for theme, row in high_prevalence.iterrows():
            prevalence = row['Prevalence_%']
            p = doc.add_paragraph()
            p.add_run(f'• {theme}: {prevalence:.1f}% prevalence')
    else:
        doc.add_paragraph('• No themes reached >70% prevalence threshold')
    
    # 2. Methods Appendix
    doc.add_heading('2. Methods Appendix', level=1)
    
    methods_p = doc.add_paragraph()
    methods_p.add_run('• Corpus: ')
    methods_p.add_run(f'{transcript_id} transcript')
    
    if alpha_report:
        overall_alpha = alpha_report.get('overall_alpha', 0)
        n_units = alpha_report.get('n_units', 0)
        methods_p = doc.add_paragraph()
        methods_p.add_run(f'• Overall Krippendorff α: {overall_alpha:.3f}')
        methods_p = doc.add_paragraph()
        methods_p.add_run(f'• Units analyzed: {n_units}')
    
    # 3. Theme Narratives
    doc.add_heading('3. Theme Narratives', level=1)
    
    # Get top 5 themes by prevalence
    top_themes = theme_matrix.head(5)
    
    for theme, row in top_themes.iterrows():
        prevalence = row['Prevalence_%']
        
        # Theme header
        theme_heading = doc.add_heading(f'#{theme}', level=2)
        
        # Prevalence
        p = doc.add_paragraph()
        p.add_run(f'Prevalence: {prevalence:.1f}%')
        
        # Exemplar quotes
        quotes = get_exemplar_quotes(transcript_id, theme)
        if quotes:
            p = doc.add_paragraph()
            p.add_run('Evidence: ')
            for i, quote in enumerate(quotes):
                if i > 0:
                    p.add_run('; ')
                p.add_run(quote)
        
        # Implication placeholder
        p = doc.add_paragraph()
        p.add_run('Implication: [Analyst to fill in]')
        
        doc.add_paragraph()  # Spacing
    
    # 4. Negative-Case Analysis
    doc.add_heading('4. Negative-Case Analysis', level=1)
    doc.add_paragraph('[Analyst to identify contradictory quotes and adjustments]')
    
    # 5. Actionable Recommendations
    doc.add_heading('5. Actionable Recommendations', level=1)
    
    # Create table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Theme'
    header_cells[1].text = 'Tool Spec'
    header_cells[2].text = 'Owner'
    header_cells[3].text = 'Priority'
    
    # Add placeholder rows for top themes
    for theme, row in top_themes.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = theme
        row_cells[1].text = '[Tool specification]'
        row_cells[2].text = '[Owner]'
        row_cells[3].text = '[Priority]'
    
    # Save document
    output_path = OUTPUT_DIR / f"{transcript_id}_topline_v1.docx"
    doc.save(output_path)
    print(f"Saved executive brief to {output_path}")


def main():
    """Generate executive briefs for all transcripts."""
    merged_files = list(OUTPUT_DIR.glob("*_tags_merged.csv"))
    
    if not merged_files:
        print(f"No merged tag files found in {OUTPUT_DIR}")
        return
    
    print(f"Found {len(merged_files)} transcript(s) for executive briefs")
    
    for merged_file in merged_files:
        transcript_id = merged_file.stem.replace("_tags_merged", "")
        create_executive_brief(transcript_id)
    
    print(f"\nAll executive briefs complete. Results in {OUTPUT_DIR}")


if __name__ == "__main__":
    main() 